#!/usr/bin/env python3
"""
Generate DOCX version of TimeToxSV_Methods_Paper.tex

Converts the LaTeX paper content into a formatted Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, font_size=None, alignment=None, space_after=None, space_before=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def add_mixed_paragraph(doc, parts, style='Normal', alignment=None, space_after=None):
    """Add paragraph with mixed formatting. parts = list of (text, bold, italic) tuples."""
    p = doc.add_paragraph(style=style)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Configure heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

    # =========================================================================
    # TITLE
    # =========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TimeToxSV: An LLM-Based Pipeline for Automated Extraction of Time Toxicity from Clinical Trial Protocols')
    run.bold = True
    run.font.size = Pt(16)
    title.paragraph_format.space_after = Pt(6)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run('[Authors -- Placeholder]\n[Affiliations -- Placeholder]')
    run.font.size = Pt(11)
    authors.paragraph_format.space_after = Pt(18)

    # =========================================================================
    # ABSTRACT
    # =========================================================================
    doc.add_heading('Abstract', level=1)

    p = doc.add_paragraph()
    p.add_run('Background: ').bold = True
    p.add_run('Time toxicity\u2014the cumulative burden of healthcare contact days imposed on patients by clinical trial participation\u2014is an increasingly recognized but understudied dimension of treatment burden. Manually extracting time toxicity from clinical trial protocols is labor-intensive, subjective, and difficult to scale.')
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run('Methods: ').bold = True
    p.add_run('We developed TimeToxSV, an LLM-based pipeline for automated extraction of healthcare contact days from Schedule of Assessments (SoA) tables in clinical trial protocols. Using Google\u2019s Gemini models, the pipeline operates in three stages: (1) summary extraction, where Gemini 2.5 Flash identifies and extracts SoA-relevant pages from full-length protocol PDFs; (2) time toxicity extraction, where Gemini 3.0 Flash quantifies healthcare contact days at six cumulative timepoints (screening, 1, 3, 6, 9, and 12 months) for each treatment arm; and (3) multi-run consensus, where position-based arm matching and median aggregation across three independent extraction runs produces stable estimates. We validated the extraction pipeline against 20 programmatically generated synthetic schedules (40 arms, 240 comparisons) with clinician-verified ground truth. We evaluated two extraction architectures\u2014a single-pass vanilla approach and a two-stage structure-then-count approach\u2014and assessed reproducibility through 5-run stability testing on 100 real-world protocols and 3-run stability testing on all 644 real-world protocols.')
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run('Results: ').bold = True
    p.add_run('On synthetic data, the two-stage pipeline achieved 100% clinical accuracy (\u00b13 days) with a mean absolute error of 0.81 days, compared to 41.5% clinical accuracy and 9.0-day MAE for the vanilla pipeline. However, on real-world protocols, the vanilla pipeline demonstrated markedly superior reproducibility: 93.7% clinical stability (IQR \u2264 3 days) across 5 runs on 100 protocols and 95.3% across 3 runs on all 644 protocols, with 82.0% achieving perfect stability (IQR = 0). The two-stage pipeline exhibited high inter-run variance on real protocols despite its synthetic accuracy. The production pipeline (vanilla + 3-run consensus) was applied to 644 real-world oncology trial protocols, successfully extracting time toxicity for 1,288 treatment arms across multiple disease sites and treatment modalities.')
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run('Conclusions: ').bold = True
    p.add_run('TimeToxSV enables scalable, reproducible quantification of time toxicity from clinical trial protocols. The pipeline demonstrates that extraction stability on real-world data, rather than accuracy on synthetic benchmarks, is the decisive factor for production deployment. Code is available at [GitHub URL].')
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run('Keywords: ').bold = True
    p.add_run('time toxicity, large language models, clinical trial protocols, healthcare contact days, schedule of assessments, automated extraction')
    p.paragraph_format.space_after = Pt(12)

    # =========================================================================
    # 1. INTRODUCTION
    # =========================================================================
    doc.add_heading('1. Introduction', level=1)

    doc.add_paragraph('Clinical trial participation imposes substantial time demands on patients beyond the direct effects of treatment itself. These demands\u2014including clinic visits, laboratory tests, imaging studies, and other protocol-mandated activities\u2014constitute what has been termed \u201ctime toxicity,\u201d defined as the cumulative number of days requiring in-person healthcare contact. Unlike traditional toxicity endpoints that focus on physical side effects, time toxicity captures the temporal burden that can affect quality of life, employment, caregiving responsibilities, and ultimately trial participation decisions.')

    doc.add_paragraph('Despite growing recognition of time toxicity as a meaningful patient-centered outcome, its systematic quantification remains challenging. The Schedule of Assessments (SoA) table\u2014a dense, multi-page document embedded within clinical trial protocols\u2014specifies all required activities, their timing, and their applicability to each treatment arm. Manually extracting healthcare contact days from these tables requires: (1) mapping cycle-based schedules to calendar days, (2) expanding grouped column headers (e.g., \u201cCycles 4\u201312\u201d) into individual cycles, (3) distinguishing arm-specific assessments from universal ones, and (4) accumulating unique contact days across cumulative time windows while avoiding double-counting. This process is time-consuming, error-prone, and difficult to standardize across reviewers.')

    doc.add_paragraph('Recent advances in large language models (LLMs) have demonstrated remarkable capabilities in processing complex structured documents, including medical texts and tabular data. LLMs can parse visual table layouts, understand domain-specific terminology, and perform multi-step reasoning over structured content\u2014skills directly applicable to SoA table interpretation. However, deploying LLMs for clinical data extraction introduces challenges of its own: output variability across runs, sensitivity to prompt design, and the difficulty of validating extractions at scale when ground truth is scarce.')

    doc.add_paragraph('In this work, we present TimeToxSV, an end-to-end LLM-based pipeline for automated extraction of time toxicity from clinical trial protocols. Our contributions include:')

    contributions = [
        'A complete pipeline from raw protocol PDFs to structured time toxicity data, using Gemini models for both document preprocessing and data extraction.',
        'A programmatically generated ground truth dataset of 20 synthetic SoA schedules with clinician-verified time toxicity values, enabling systematic validation.',
        'A comparative evaluation of two extraction architectures\u2014single-pass (vanilla) and two-stage (structure-then-count)\u2014revealing a critical trade-off between synthetic accuracy and real-world stability.',
        'A position-based consensus mechanism that handles LLM-generated arm name instability across runs, enabling robust multi-run aggregation.',
        'Application of the production pipeline to 644 real-world oncology trial protocols, demonstrating feasibility at scale.',
    ]
    for i, item in enumerate(contributions, 1):
        doc.add_paragraph(f'{i}. {item}', style='List Number')

    # =========================================================================
    # 2. DATA COLLECTION AND PREPROCESSING
    # =========================================================================
    doc.add_heading('2. Data Collection and Preprocessing', level=1)

    doc.add_heading('2.1 Protocol Acquisition', level=2)
    doc.add_paragraph('We compiled a dataset of 649 clinical trial protocol documents from ClinicalTrials.gov, spanning multiple oncology disease sites including breast, thoracic, gastrointestinal, genitourinary, gynecologic, head and neck, skin, central nervous system, and hematologic malignancies. Protocols were selected based on the availability of full-text documents with Schedule of Assessments tables.')

    doc.add_heading('2.2 Summary Extraction', level=2)
    doc.add_paragraph('Full-length clinical trial protocols typically span 100\u2013400 pages, of which only 5\u201315 pages contain the SoA table and related scheduling information. To reduce input size and improve extraction accuracy, we developed an automated summary extraction stage using Google\u2019s Gemini 2.5 Flash model.')

    doc.add_paragraph('For each protocol PDF, the summary extraction pipeline:')
    steps = [
        'Uploads the full PDF to the Gemini API.',
        'Prompts the model to identify pages containing SoA tables, visit schedules, assessment timing, or related scheduling information. The prompt instructs the model to be liberal in page identification to avoid missing relevant content.',
        'Generates a consolidated summary PDF containing: (a) a concise trial summary (study design, treatment arms, duration), and (b) the identified SoA pages plus one surrounding page on each side for context.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph('The summary extraction uses forced JSON output via response_mime_type="application/json" with a structured schema, temperature 0.0, and top-p 0.95. Processing time averages 2\u20133 minutes per protocol. Of the 649 protocols processed, summaries were successfully generated for 646, yielding a 99.5% success rate.')

    # Table 1: Dataset Overview
    p = doc.add_paragraph()
    p.add_run('Table 1: Dataset Overview').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(doc,
        ['Characteristic', 'Value'],
        [
            ['Total protocols collected', '649'],
            ['Protocols with successful summary extraction', '644'],
            ['Summary extraction success rate', '99.5%'],
            ['Summary extraction model', 'Gemini 2.5 Flash'],
            ['Average processing time per protocol', '2\u20133 min'],
            ['Average SoA pages identified per protocol', '5\u201315'],
            ['Disease sites represented', '11'],
        ])

    doc.add_paragraph()  # spacer

    # =========================================================================
    # 3. GROUND TRUTH DEVELOPMENT
    # =========================================================================
    doc.add_heading('3. Ground Truth Development', level=1)

    doc.add_heading('3.1 Synthetic Schedule Generation', level=2)
    doc.add_paragraph('To enable systematic validation of extraction pipelines, we developed a programmatic generator for synthetic SoA schedules. The generator (generate_schedules.py, 1,284 lines) creates realistic HTML-rendered schedules with known ground truth values, enabling precise accuracy measurement without the ambiguity inherent in real-world protocol interpretation.')

    doc.add_paragraph('Each synthetic schedule includes:')
    items = [
        'A complete SoA table with realistic assessment categories (administrative, clinical, laboratory, cardiac, safety, treatment, imaging, and patient-reported outcomes).',
        'Two treatment arms (intervention and control) with distinct visit schedules.',
        'Cycle-based timing with configurable cycle lengths (7, 21, 28, or 35 days).',
        'Treatment durations ranging from 2 to 12 months.',
        'A legend section specifying arm-specific visit patterns.',
        'Footnotes, headers, and formatting consistent with real clinical trial protocols.',
        'Disease-specific assessments appropriate to each tumor type.',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('We generated 20 synthetic schedules spanning 8 oncology disease categories (breast, thoracic, gastrointestinal, genitourinary, head & neck, melanoma, gynecologic, and sarcoma) with three complexity levels: simple (n=5, single visit day per cycle), moderate (n=10, two visit days per cycle), and complex (n=5, three visit days per cycle). Five distinct visual styles were used to test robustness to formatting variation. Schedules incorporated diverse treatment modalities including systemic therapy, radiation, and surgery.')

    doc.add_heading('3.2 Ground Truth Calculation', level=2)
    doc.add_paragraph('Ground truth healthcare contact days were computed deterministically from each schedule\u2019s configuration parameters:')
    calc_steps = [
        'Treatment visit days were mapped from cycle-based schedules to absolute calendar days.',
        'Imaging days were added at protocol-specified intervals.',
        'End-of-treatment (EOT) and follow-up visits at 9 and 12 months were included.',
        'Unique calendar days were counted at each cumulative timepoint (screening, 1, 3, 6, 9, 12 months).',
    ]
    for i, step in enumerate(calc_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph('This yielded 40 arms \u00d7 6 timepoints = 240 ground truth comparisons.')

    # Table 2: Synthetic Schedule Characteristics
    p = doc.add_paragraph()
    p.add_run('Table 2: Synthetic Schedule Characteristics').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(doc,
        ['Characteristic', 'Value'],
        [
            ['Total schedules', '20'],
            ['Total arms', '40 (20 intervention, 20 control)'],
            ['Total ground truth comparisons', '240'],
            ['Disease categories', '8'],
            ['Complexity: Simple / Moderate / Complex', '5 / 10 / 5'],
            ['Cycle lengths', '7, 21, 28, 35 days'],
            ['Treatment durations', '2\u201312 months'],
            ['Visual style variants', '5'],
            ['Treatment modalities', 'Systemic, Radiation, Surgery'],
        ])

    doc.add_paragraph()

    doc.add_heading('3.3 Ground Truth Verification', level=2)
    doc.add_paragraph('The programmatically generated ground truth values were independently verified by multiple reviewers, including a clinician expert who assessed the clinical plausibility of the synthetic schedules and confirmed the correctness of the calculated contact days. This verification ensured that both the schedule designs and the ground truth calculations accurately reflected real-world clinical trial visit patterns.')

    # =========================================================================
    # 4. LLM PIPELINE DESIGN
    # =========================================================================
    doc.add_heading('4. LLM Pipeline Design', level=1)

    doc.add_paragraph('We evaluated two extraction architectures, both using Google\u2019s Gemini 3.0 Flash (gemini-3-flash-preview) as the base model with temperature 0.1 to balance accuracy and determinism.')

    doc.add_heading('4.1 Vanilla Pipeline (Primary)', level=2)
    doc.add_paragraph('The vanilla pipeline performs single-pass extraction, providing the model with a summary SoA PDF and a structured prompt that requests direct quantification of healthcare contact days. The prompt (76 lines of instruction, excluding the JSON output template) specifies:')

    prompt_parts = [
        ('Definitions: ', True, False),
        ('Healthcare contact day definition, six cumulative time windows (screening through 12 months), and the instruction to count each unique calendar day once regardless of the number of assessments.', False, False),
    ]
    add_mixed_paragraph(doc, prompt_parts)

    prompt_parts2 = [
        ('Extraction rules: ', True, False),
        ('Five key rules covering arm identification, cycle-to-calendar mapping, grouped column expansion, unique day counting, and arm-specific visit handling (e.g., superscript markers like X', False, False),
        ('A', False, False),
        (').', False, False),
    ]
    add_mixed_paragraph(doc, prompt_parts2)

    prompt_parts3 = [
        ('Category definitions: ', True, False),
        ('Four categories of healthcare contact\u2014core treatment, imaging/diagnostics, labs, and clinic visits\u2014with the note that a single day may span multiple categories.', False, False),
    ]
    add_mixed_paragraph(doc, prompt_parts3)

    prompt_parts4 = [
        ('Output format: ', True, False),
        ('Strict JSON schema specifying arm name, intervention type, healthcare contact days at each timepoint, category breakdown, and extraction notes (cycle length, treatment duration, visit pattern, disease).', False, False),
    ]
    add_mixed_paragraph(doc, prompt_parts4)

    doc.add_paragraph('The model receives the PDF as a visual input and produces a JSON response. Three retry attempts with exponential backoff handle transient API failures.')

    # Code listing
    p = doc.add_paragraph()
    p.add_run('Listing 1: Condensed excerpt of the vanilla extraction prompt.').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    code_text = '''You are an expert clinical trial protocol analyst.
Your task is to extract healthcare contact days
from the Schedule of Assessments (SoA) table.

## DEFINITIONS
**Healthcare Contact Day**: Any calendar day requiring
in-person contact with the healthcare system.
Count each unique calendar day ONCE.

**Time Windows** (cumulative):
- screening: All pre-treatment visits
- 1_month: Day 1 through Day 30
- 3_months: Day 1 through Day 90
...
- 12_months: Day 1 through Day 365

## EXTRACTION RULES
1. Identify all treatment arms
2. Map cycle structure to calendar days
3. CRITICAL: Expand grouped columns (e.g., "C4-C12")
4. Count unique visit days
5. Handle arm-specific visits (superscripts)

## OUTPUT FORMAT
Return ONLY valid JSON: [{"arm_name": "...",
  "healthcare_contact_days": {...}, ...}]'''

    code_p = doc.add_paragraph()
    code_run = code_p.add_run(code_text)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)
    code_p.paragraph_format.space_before = Pt(6)
    code_p.paragraph_format.space_after = Pt(6)

    doc.add_heading('4.2 Two-Stage Pipeline (Alternative)', level=2)
    doc.add_paragraph('The two-stage pipeline decomposes extraction into two sequential LLM calls:')

    p = doc.add_paragraph()
    p.add_run('Stage 1 \u2014 Structure Extraction: ').bold = True
    p.add_run('The model extracts the structural parameters of the SoA table, including: cycle length in days, visit days per cycle for each arm, treatment duration, special visit information (screening days, EOT visit, follow-up visits), and disease type. The output is a structured JSON object capturing the schedule\u2019s \u201cblueprint\u201d without performing any counting. This stage uses forced JSON output (response_mime_type="application/json").')

    p = doc.add_paragraph()
    p.add_run('Stage 2 \u2014 Count Calculation: ').bold = True
    p.add_run('The extracted structure from Stage 1, formatted as JSON, is provided alongside the original PDF as context for a second LLM call. This call performs the arithmetic: mapping cycles to calendar days, applying visit patterns, and accumulating unique contact days at each timepoint. The prompt includes explicit calculation rules (e.g., \u201cTotal cycles = floor(treatment_duration_months \u00d7 30 / cycle_length_days)\u201d).')

    doc.add_paragraph('The rationale for decomposition was to separate the perceptual task (reading the table) from the computational task (counting days), hypothesizing that this division would reduce errors. As described in Section 6, while this approach achieved superior accuracy on synthetic data, it exhibited unacceptable variability on real-world protocols.')

    # =========================================================================
    # 5. VALIDATION ON SYNTHETIC DATA
    # =========================================================================
    doc.add_heading('5. Validation on Synthetic Data', level=1)

    doc.add_heading('5.1 Evaluation Metrics', level=2)
    doc.add_paragraph('We evaluated both pipelines using three metrics computed across all 240 ground truth comparisons (20 schedules \u00d7 2 arms \u00d7 6 timepoints):')

    metrics = [
        'Exact Match Accuracy: Percentage of extractions matching the ground truth exactly.',
        'Clinical Accuracy (\u00b13 days): Percentage of extractions within 3 days of the ground truth, reflecting the threshold at which differences are clinically meaningful.',
        'Mean Absolute Error (MAE): Average absolute difference between extracted and ground truth values, in days.',
    ]
    for m in metrics:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_heading('5.2 Results', level=2)
    doc.add_paragraph('The two-stage pipeline substantially outperformed the vanilla pipeline on all synthetic accuracy metrics, achieving perfect clinical accuracy. However, as Section 6 demonstrates, these synthetic metrics did not predict real-world performance.')

    # Table 3: Validation Results
    p = doc.add_paragraph()
    p.add_run('Table 3: Validation Results on 20 Synthetic Schedules (240 Comparisons)').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(doc,
        ['Pipeline', 'Exact Match', 'Clinical Accuracy (\u00b13d)', 'MAE (days)'],
        [
            ['Vanilla (single-pass)', '0.3%', '41.5%', '9.0'],
            ['Two-Stage (structure \u2192 count)', '29.2%', '100.0%', '0.81'],
        ])

    doc.add_paragraph()
    doc.add_paragraph('The vanilla pipeline\u2019s low synthetic accuracy is attributable to the inherent difficulty of performing multi-step arithmetic (cycle mapping, grouped column expansion, cumulative counting) in a single LLM call. The model frequently miscounts visits in grouped columns or makes errors in cycle-to-day conversion. In contrast, the two-stage pipeline benefits from explicit structural extraction followed by guided calculation, reducing the cognitive load on each individual call.')

    # =========================================================================
    # 6. STABILITY AND REPRODUCIBILITY TESTING
    # =========================================================================
    doc.add_heading('6. Stability and Reproducibility Testing', level=1)

    doc.add_heading('6.1 Motivation', level=2)
    doc.add_paragraph('While synthetic accuracy measures correctness against known ground truth, production deployment requires a different property: reproducibility. If the same protocol yields substantially different time toxicity values across independent extraction runs, the resulting dataset cannot support reliable downstream analysis. We therefore conducted systematic stability testing on real-world protocols.')

    doc.add_heading('6.2 Phase 1: Pilot Stability Study (5 Runs, 100 Protocols)', level=2)
    doc.add_paragraph('We first compared both pipelines on a pilot set of real-world protocols. Using 5 CTGProtocol summary PDFs, we ran each pipeline 5 times and measured the inter-quartile range (IQR) of the 12-month healthcare contact day count across runs for each arm.')

    doc.add_paragraph('The two-stage pipeline exhibited substantial inter-run variability on real protocols. For example, one protocol arm produced values of [94, 473, 473, 540, 543] across 5 two-stage runs (IQR = 258.0), while the same arm under the vanilla pipeline produced [99, 105, 106, 106, 106] (IQR = 4.0). Another protocol yielded two-stage values of [27, 53, 53, 53, 53] (IQR = 13.0) versus vanilla values of [8, 8, 8, 8, 8] (IQR = 0.0).')

    doc.add_paragraph('This discrepancy\u2014high synthetic accuracy but high real-world variability for the two-stage pipeline\u2014reflects a fundamental challenge: real-world SoA tables are far more heterogeneous than synthetic schedules. Ambiguous column headers, inconsistent formatting, multi-page tables, and complex footnotes create interpretation variability that is amplified by the two-stage pipeline\u2019s reliance on accurate structural extraction in Stage 1.')

    doc.add_paragraph('We then extended stability testing to 100 real-world protocols with 5 independent vanilla runs.')

    # Table 4: Stability Phase 1
    p = doc.add_paragraph()
    p.add_run('Table 4: Vanilla Pipeline Stability: 5 Runs on 100 Real-World Protocols').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(doc,
        ['Stability Category', 'Arms', 'Percentage'],
        [
            ['Perfect Stability (IQR = 0)', '417', '77.2%'],
            ['Clinical Stability (IQR \u2264 3)', '506', '93.7%'],
            ['High Variance (IQR > 3)', '34', '6.3%'],
            ['Total arms analyzed', '540', ''],
        ])

    doc.add_paragraph()

    doc.add_heading('6.3 Phase 2: Full-Scale Stability (3 Runs, 644 Protocols)', level=2)
    doc.add_paragraph('Based on the pilot results, we selected the vanilla pipeline for production deployment and ran 3 independent extractions on all 644 successfully preprocessed protocols.')

    # Table 5: Stability Phase 2
    p = doc.add_paragraph()
    p.add_run('Table 5: Vanilla Pipeline Stability: 3 Runs on 644 Real-World Protocols').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(doc,
        ['Stability Category', 'Arms', 'Percentage'],
        [
            ['Perfect Stability (IQR = 0)', '2,172', '82.0%'],
            ['Clinical Stability (IQR \u2264 3)', '2,524', '95.3%'],
            ['High Variance (IQR > 3)', '125', '4.7%'],
            ['Total arms analyzed', '2,649', ''],
        ])

    doc.add_paragraph()
    doc.add_paragraph('The improvement from 93.7% to 95.3% clinical stability when moving from 5 runs on 100 protocols to 3 runs on 644 protocols likely reflects both the natural variance reduction with fewer runs and the broader protocol sample smoothing out edge cases.')

    doc.add_heading('6.4 Stability Metrics', level=2)
    doc.add_paragraph('We define clinical stability as follows:')
    stab_defs = [
        'Perfect Stability (IQR = 0): All runs produced identical 12-month contact day counts for this arm.',
        'Clinical Stability (IQR \u2264 3): The interquartile range of 12-month counts is at most 3 days\u2014a threshold below which differences are not clinically meaningful.',
        'High Variance (IQR > 3): Substantial disagreement across runs, flagging arms for potential review.',
    ]
    for d in stab_defs:
        doc.add_paragraph(d, style='List Bullet')

    # =========================================================================
    # 7. PRODUCTION PIPELINE AND CONSENSUS
    # =========================================================================
    doc.add_heading('7. Production Pipeline and Consensus', level=1)

    doc.add_heading('7.1 Pipeline Architecture', level=2)
    doc.add_paragraph('Based on the validation and stability results, we adopted the vanilla pipeline with 3-run consensus as the production configuration.')

    # Pipeline flow as text (since we can't render TikZ)
    p = doc.add_paragraph()
    p.add_run('Figure 1: TimeToxSV Production Pipeline Architecture').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    flow = doc.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow_text = (
        'Full Protocol PDFs (649 protocols)\n'
        '        \u2193\n'
        'Summary Extraction (Gemini 2.5 Flash)\n'
        '    [Identifies SoA pages, generates consolidated PDF]\n'
        '        \u2193\n'
        'SoA Summary PDFs (644 protocols)\n'
        '        \u2193\n'
        'Vanilla Extraction \u00d7 3 Runs (Gemini 3.0 Flash)\n'
        '    [Single-pass extraction, temp=0.1, JSON output]\n'
        '        \u2193\n'
        'Position-Based Consensus (Median Aggregation)\n'
        '    [Matches arms by position, takes median across runs]\n'
        '        \u2193\n'
        'Final Dataset (644 protocols, 1,288 arms)'
    )
    run = flow.add_run(flow_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading('7.2 Position-Based Consensus Mechanism', level=2)
    doc.add_paragraph('A key challenge in multi-run LLM extraction is arm name instability: the same treatment arm may be labeled differently across runs (e.g., \u201cArm A (Pembrolizumab + Chemotherapy)\u201d vs. \u201cPembrolizumab + Chemotherapy Arm\u201d vs. \u201cTreatment Arm A: Pembro/Chemo\u201d). Standard name-based matching fails in this scenario.')

    doc.add_paragraph('We address this through position-based matching:')
    matching_steps = [
        'Within each run, arms are grouped by (filename, intervention_type).',
        'Within each group, arms are sorted by their 12-month healthcare contact day count.',
        'A positional index (pos_idx) is assigned based on this sorted order.',
        'Consensus is computed by grouping across runs by (filename, intervention_type, pos_idx) and taking the median of each timepoint.',
    ]
    for i, step in enumerate(matching_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph('This approach leverages the insight that even when arm names vary, the relative ordering of arms by time toxicity burden is stable. The intervention type (intervention vs. control) is normalized during preprocessing (e.g., \u201cplacebo\u201d \u2192 \u201ccontrol\u201d, \u201cactive_comparator\u201d \u2192 \u201cintervention\u201d).')

    # Table 6: Consensus Stats
    p = doc.add_paragraph()
    p.add_run('Table 6: Production Consensus Dataset Statistics').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(doc,
        ['Metric', 'Value'],
        [
            ['Protocols processed', '644'],
            ['Total consensus arms', '1,288'],
            ['Arms with 3 concordant runs', 'majority'],
            ['Extraction model', 'Gemini 3.0 Flash'],
            ['Extraction temperature', '0.1'],
            ['Extraction runs per protocol', '3'],
            ['Consensus method', 'Median aggregation'],
            ['Average processing time per extraction', '~4 min'],
        ])

    doc.add_paragraph()

    doc.add_heading('7.3 Cost and Runtime', level=2)
    doc.add_paragraph('The full production pipeline cost is dominated by API calls to the Gemini models. Summary extraction (Gemini 2.5 Flash) processes each protocol in approximately 2\u20133 minutes. Each vanilla extraction run (Gemini 3.0 Flash) takes approximately 4 minutes per protocol. With 3 runs across 644 protocols, the total extraction time is approximately 128 hours of sequential processing (reducible through parallelization). The Gemini API pricing makes the total cost well under $100 for the complete dataset, making the approach economically viable for large-scale studies.')

    # =========================================================================
    # 8. APPLICATION
    # =========================================================================
    doc.add_heading('8. Application', level=1)

    doc.add_paragraph('To demonstrate the pipeline\u2019s utility, we applied the production consensus dataset to a landscape analysis of time toxicity across oncology clinical trials. The full clinical analysis\u2014including descriptive statistics, comparative analyses across disease sites and treatment modalities, temporal trends, multivariable regression, and component analysis\u2014is presented in a companion paper.')

    doc.add_paragraph('Briefly, the pipeline enabled analysis of time toxicity patterns across 644 protocols spanning 11 disease sites, multiple treatment modalities (chemotherapy, immunotherapy, targeted therapy, endocrine therapy, radiation, and procedures), and trial start years from 1994 to 2021. The median 12-month time toxicity across all intervention arms was quantified, and systematic differences by disease site, treatment type, and funding source were characterized\u2014analyses that would be infeasible through manual extraction at this scale.')

    # =========================================================================
    # 9. DISCUSSION
    # =========================================================================
    doc.add_heading('9. Discussion', level=1)

    doc.add_heading('9.1 Accuracy vs. Stability Trade-off', level=2)
    doc.add_paragraph('The central finding of this work is the critical distinction between synthetic accuracy and real-world stability. The two-stage pipeline achieved perfect clinical accuracy on synthetic schedules but exhibited unacceptable variability on real-world protocols. Conversely, the vanilla pipeline showed modest synthetic accuracy but remarkable stability\u201495.3% of arms had an IQR \u2264 3 days across independent runs.')

    doc.add_paragraph('This divergence has a straightforward explanation. Synthetic schedules have unambiguous structure: clear column headers, explicit legends, and well-defined visit patterns. The two-stage pipeline excels when the structural extraction step can cleanly parse this structure. Real-world protocols, however, contain multi-page tables with complex footnotes, ambiguous column groupings, inconsistent notation, and formatting artifacts. In this environment, small variations in structural extraction at Stage 1 cascade into large counting differences at Stage 2. The vanilla pipeline, by avoiding explicit structural decomposition, is more robust to these sources of ambiguity\u2014it treats extraction as a holistic task, reducing the number of failure points.')

    doc.add_heading('9.2 Design Decisions', level=2)

    p = doc.add_paragraph()
    p.add_run('Temperature selection: ').bold = True
    p.add_run('We used temperature 0.1 throughout extraction. Lower temperatures (0.0) with forced JSON output were tested during development but did not consistently improve stability, while they occasionally caused the model to produce malformed outputs. Temperature 0.1 provided a balance between determinism and output quality.')

    p = doc.add_paragraph()
    p.add_run('Prompt design: ').bold = True
    p.add_run('The vanilla prompt underwent iterative refinement. Key additions included the explicit instruction to expand grouped columns (\u201cCRITICAL \u2014 Expand grouped columns\u201d), the specification that time windows are cumulative, and the category breakdown requirement. Each addition addressed a systematic error pattern observed during development.')

    p = doc.add_paragraph()
    p.add_run('Three-run consensus: ').bold = True
    p.add_run('Three runs were chosen as a pragmatic balance between stability and cost. Increasing from 3 to 5 runs would provide marginal stability improvement (as evidenced by the 93.7% vs. 95.3% comparison) at a 67% increase in cost.')

    doc.add_heading('9.3 Limitations', level=2)
    limitations = [
        'Single LLM family: All extraction uses Google\u2019s Gemini models. While Gemini 3.0 Flash demonstrated good performance, comparative evaluation against other LLM families (GPT-4, Claude) was not conducted with formal benchmarking.',
        'Synthetic-only formal validation: Clinical accuracy was formally validated only against synthetic schedules. Real-world validation relied on stability testing and informal spot-checking against manually reviewed protocols rather than a comprehensive gold standard.',
        'Systematic bias: The vanilla pipeline may systematically under- or over-count healthcare contact days. While the consensus mechanism reduces random error, it cannot correct consistent directional bias across runs.',
        'Protocol heterogeneity: Extremely complex protocols (e.g., multi-phase adaptive designs with dose-escalation cohorts) may exceed the pipeline\u2019s extraction capability, contributing to the 4.7% of arms with high variance.',
        'Temporal coverage: The pipeline extracts contact days up to 12 months. Longer-term follow-up visits beyond 12 months are not captured.',
    ]
    for i, lim in enumerate(limitations, 1):
        doc.add_paragraph(f'{i}. {lim}', style='List Number')

    doc.add_heading('9.4 Comparison to Related Work', level=2)
    doc.add_paragraph('LLM-based information extraction from medical documents is an active area of research. Prior work has demonstrated the feasibility of using LLMs for tasks such as extracting structured data from clinical notes, annotating medical literature, and processing clinical trial registries. Our work extends this line of research to a specific, previously unautomated task\u2014time toxicity quantification from SoA tables\u2014that requires both visual table parsing and multi-step arithmetic reasoning.')

    doc.add_paragraph('The stability-focused evaluation framework we employ is less common in the LLM extraction literature, which typically emphasizes accuracy against gold standards. Our finding that stability on heterogeneous real-world data is a better predictor of production utility than accuracy on controlled benchmarks may generalize to other LLM-based extraction tasks.')

    # =========================================================================
    # 10. CONCLUSION
    # =========================================================================
    doc.add_heading('10. Conclusion', level=1)

    doc.add_paragraph('We presented TimeToxSV, an LLM-based pipeline for automated extraction of time toxicity from clinical trial protocol SoA tables. The pipeline combines Gemini-based summary extraction, vanilla single-pass time toxicity extraction, and position-based multi-run consensus to produce stable, scalable quantification of healthcare contact days.')

    doc.add_paragraph('Our evaluation revealed a critical insight for LLM deployment in clinical data extraction: performance on controlled synthetic benchmarks does not predict real-world reliability. The two-stage pipeline\u2019s 100% synthetic accuracy was undermined by high variability on real protocols, while the vanilla pipeline\u2019s modest 41.5% synthetic accuracy translated to 95.3% clinical stability in production\u2014the metric that ultimately matters for downstream analysis.')

    doc.add_paragraph('The production pipeline was successfully applied to 644 real-world oncology trial protocols, enabling large-scale quantification of time toxicity that would be infeasible through manual extraction. Future work includes expanding ground truth validation with manually annotated real-world protocols, evaluating multi-model ensemble approaches, and extending the framework to other structured clinical trial data extraction tasks.')

    doc.add_heading('Code Availability', level=2)
    doc.add_paragraph('The complete pipeline code, synthetic schedule generator, and analysis scripts are available at [GitHub URL].')

    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_heading('References', level=1)

    references = [
        '[1] Gupta A, Eisenhauer EA, Booth CM. The time toxicity of cancer treatment. Journal of Clinical Oncology. 2022;40(15):1611\u20131615.',
        '[2] Shah KP, et al. Time toxicity: a framework for assessing the patient burden of treatment. Journal of the National Cancer Institute. 2024.',
        '[3] Google DeepMind. Gemini: A family of highly capable multimodal models. arXiv preprint. 2024.',
        '[4] Nori H, et al. Capabilities of GPT-4 on medical challenge problems. arXiv preprint arXiv:2303.13375. 2023.',
        '[5] Singhal K, et al. Large language models encode clinical knowledge. Nature. 2023;620(7972):172\u2013180.',
        '[6] Jin Q, et al. Matching patients to clinical trials with large language models. arXiv preprint arXiv:2307.15051. 2024.',
        '[7] Agrawal M, Hegselmann S, Lang H, Kim Y, Sontag D. Large language models are few-shot clinical information extractors. arXiv preprint arXiv:2205.12689. 2022.',
        '[8] Tang L, et al. Evaluating large language models on medical evidence summarization. npj Digital Medicine. 2023;6(1):158.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)

    # =========================================================================
    # SAVE
    # =========================================================================
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TimeToxSV_Methods_Paper.docx')
    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")


if __name__ == '__main__':
    main()
