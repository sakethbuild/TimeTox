#!/usr/bin/env python3
"""
Generate DOCX version of JAMA_Oncology_Brief_Report.tex
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def setup_styles(doc):
    """Configure document styles for JAMA format."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 2.0

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Times New Roman'
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.font.bold = True
        if level == 1:
            h.font.size = Pt(14)
        elif level == 2:
            h.font.size = Pt(12)
        else:
            h.font.size = Pt(11)


def add_mixed(doc, parts, style='Normal', alignment=None, space_after=None):
    """Add paragraph with mixed bold/italic runs. parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph(style=style)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_para(doc, text, bold=False, italic=False, font_size=None, alignment=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            if isinstance(val, str) and val.startswith('\t'):
                run.italic = True

    return table


def main():
    doc = Document()
    setup_styles(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ========================================================================
    # TITLE
    # ========================================================================
    add_para(doc, "Time Toxicity in Phase III Oncology Clinical Trials:\nA Landscape Analysis of 644 Protocols",
             bold=True, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "[Author List]\n[Affiliations]",
             font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ========================================================================
    # KEY POINTS BOX
    # ========================================================================
    add_para(doc, "KEY POINTS", bold=True, font_size=11, space_after=4)
    add_mixed(doc, [
        ("Question: ", True, False),
        ("What is the magnitude of time toxicity across contemporary phase III oncology clinical trials?", False, False)
    ], space_after=4)
    add_mixed(doc, [
        ("Findings: ", True, False),
        ("In this cross-sectional analysis of 644 phase III oncology protocols, median time toxicity was 19.0 healthcare contact days at 12 months (1.6 days per month), comparable to routine off-protocol cancer care. Over one-quarter (26.4%) of this burden was attributable to research-specific monitoring.", False, False)
    ], space_after=4)
    add_mixed(doc, [
        ("Meaning: ", True, False),
        ("Time toxicity in oncology trials is lower than commonly assumed, but the modifiable research-specific component presents a concrete target for protocol optimization through decentralized trial designs.", False, False)
    ], space_after=12)

    # ========================================================================
    # STRUCTURED ABSTRACT
    # ========================================================================
    doc.add_heading('Abstract', level=1)

    add_mixed(doc, [
        ("Importance: ", True, False),
        ("Time toxicity\u2014the cumulative healthcare contact days imposed on patients during clinical trial participation\u2014is increasingly recognized as a patient-centered outcome, yet its magnitude across the oncology trial landscape has never been systematically quantified.", False, False)
    ], space_after=4)

    add_mixed(doc, [
        ("Objective: ", True, False),
        ("To create a landmark reference dataset quantifying time toxicity across phase III oncology protocols, stratified by disease site and treatment modality, using a validated large language model (LLM) extraction pipeline.", False, False)
    ], space_after=4)

    add_mixed(doc, [
        ("Design, Setting, and Participants: ", True, False),
        ("Cross-sectional analysis of 644 phase III oncology trial protocols with trial start years spanning 1994 to 2021, identified from ClinicalTrials.gov. The dataset comprises 1466 treatment arms (790 intervention, 676 control) across 11 disease sites and 6 treatment modalities. Healthcare contact days were extracted from Schedules of Assessments using a validated LLM pipeline (Gemini 3.0 Flash) with three-run median consensus achieving 95.3% clinical stability.", False, False)
    ], space_after=4)

    add_mixed(doc, [
        ("Main Outcomes and Measures: ", True, False),
        ("Total healthcare contact days at 12 months, time toxicity intensity score (days per month), incremental time toxicity (intervention minus control), and component decomposition into core treatment versus research-specific burden.", False, False)
    ], space_after=4)

    add_mixed(doc, [
        ("Results: ", True, False),
        ("Median 12-month time toxicity was 19.0 days (IQR, 14.0\u201329.0) for intervention arms, corresponding to an intensity of 1.6 days per month (IQR, 1.2\u20132.4)\u2014comparable to routine off-protocol cancer care. The median incremental time toxicity was 0.0 days, with 67.6% of trials imposing identical scheduling on both arms. Time toxicity varied significantly by disease site (Kruskal-Wallis H = 59.0, P < .001), ranging from 41.0 days (CNS) to 16.0 days (breast), and by treatment modality (H = 105.6, P < .001), from 32.0 days (chemotherapy) to 11.0 days (other/supportive). Research-specific monitoring (imaging and laboratory testing) accounted for 26.4% of total time toxicity.", False, False)
    ], space_after=4)

    add_mixed(doc, [
        ("Conclusions and Relevance: ", True, False),
        ("In this landmark analysis of 644 phase III oncology protocols, time toxicity at a median intensity of 1.6 days per month is lower than commonly assumed and comparable to standard cancer care. However, over one-quarter of this burden is attributable to research-specific monitoring\u2014a modifiable component that could be reduced through decentralized trial designs. This 644-protocol dataset and validated extraction tool provide benchmark references for prospective protocol evaluation and optimization.", False, False)
    ], space_after=8)

    add_mixed(doc, [("Study Type: ", True, False), ("Cross-Sectional Study", False, False)], space_after=4)
    add_mixed(doc, [("Data Sharing Statement: ", True, False), ("The consensus dataset and extraction pipeline code are available upon request from the corresponding author.", False, False)], space_after=12)

    # ========================================================================
    # INTRODUCTION
    # ========================================================================
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        "Time toxicity\u2014defined as the cumulative healthcare contact days required by clinical trial "
        "participation\u2014has emerged as a critical patient-centered outcome. Unlike traditional adverse "
        "events that affect a subset of participants, time toxicity is borne by every enrolled patient, "
        "regardless of treatment efficacy or disease response. For patients with advanced cancer and "
        "limited life expectancy, healthcare contact days represent a significant and irreversible "
        "fraction of remaining life, making their quantification an ethical imperative."
    )

    doc.add_paragraph(
        "Despite growing recognition of this burden, no comprehensive reference dataset exists to "
        "quantify time toxicity across the oncology trial landscape. Prior studies have characterized "
        "time toxicity in selected trials or specific disease settings, but the labor-intensive nature "
        "of manually extracting visit schedules from protocol documents has precluded large-scale "
        "analysis. We sought to address this gap by leveraging a validated large language model (LLM) "
        "extraction pipeline to create a landmark reference dataset of 644 phase III oncology protocols. "
        "This analysis provides both a comprehensive characterization of time toxicity by disease site "
        "and treatment modality and a scalable tool for prospective protocol assessment."
    )

    # ========================================================================
    # METHODS
    # ========================================================================
    doc.add_heading('Methods', level=1)

    doc.add_heading('Protocol Identification and Metadata', level=2)
    doc.add_paragraph(
        "We identified 649 phase III oncology trial protocols from ClinicalTrials.gov (trial start years, "
        "1994\u20132021) with available full-text protocol documents containing Schedules of Assessments. "
        "Five protocols were excluded owing to automated processing failures (3 at summary generation, "
        "2 at extraction), yielding 644 protocols (99.2%) for analysis. Each protocol was enriched with "
        "metadata including disease site (11 categories), treatment modality (6 categories), sponsorship "
        "(industry vs nonindustry), and trial start year using structured data from ClinicalTrials.gov."
    )

    doc.add_heading('LLM Extraction Pipeline', level=2)
    doc.add_paragraph(
        "Healthcare contact days were extracted using a two-stage automated pipeline validated for this "
        "purpose. In stage 1, a large language model (Gemini 2.5 Flash; Google) identified Schedule of "
        "Assessment pages from full-length protocol PDFs (mean document length, approximately 200 pages) "
        "and generated consolidated summary documents containing only the relevant tabular content. In "
        "stage 2, a second model (Gemini 3.0 Flash; Google) performed single-pass extraction using a "
        "structured prompt that maps cycle-based visit schedules (eg, 21-day or 28-day cycles) to "
        "calendar days and quantifies per-arm healthcare contact days at 6 cumulative timepoints "
        "(screening through 12 months) with decomposition into 4 component categories: core treatment, "
        "imaging and diagnostics, laboratory testing, and clinic visits. A healthcare contact day was "
        "defined as any calendar day requiring in-person contact with the healthcare system, counted "
        "once regardless of the number of assessments on that day."
    )

    doc.add_paragraph(
        "To ensure reproducibility, three independent extraction runs were performed on all 644 "
        "protocols. Arms were matched across runs using position-based alignment within protocol and "
        "intervention-type groups, and consensus values were derived via median aggregation. The "
        "pipeline achieved 95.3% clinical stability (interquartile range [IQR] \u2264 3 days across "
        "runs) and 82.0% perfect reproducibility (IQR = 0). Validation against 20 synthetic schedules "
        "with clinician-verified ground truth demonstrated high concordance."
    )

    doc.add_heading('Outcomes', level=2)
    doc.add_paragraph(
        "The primary outcome was total healthcare contact days at 12 months per arm. Secondary outcomes "
        "included the intensity score (12-month time toxicity divided by 12, reported as days per month), "
        "incremental time toxicity (delta; intervention arm minus control arm), and component "
        "decomposition into core treatment versus research-specific burden (imaging plus laboratory monitoring)."
    )

    doc.add_heading('Statistical Analysis', level=2)
    doc.add_paragraph(
        "This analysis is intentionally descriptive: the primary contribution is the benchmark dataset "
        "itself rather than inferential modeling. We report medians with IQRs for continuous variables "
        "and counts with percentages for categorical variables. Between-group comparisons used the "
        "Kruskal-Wallis H test, with Bonferroni-corrected pairwise Wilcoxon rank-sum tests for post "
        "hoc analysis. Analyses were performed in Python 3.9 using scipy and pandas."
    )

    # ========================================================================
    # RESULTS
    # ========================================================================
    doc.add_heading('Results', level=1)

    doc.add_paragraph(
        "The final cohort comprised 644 phase III oncology trials with 1466 consensus arms (790 "
        "intervention, 676 control), including 638 matched intervention-control pairs (99.1%). "
        "Sponsorship was predominantly industry (504 [78.3%]) with median trial start year of 2013 "
        "(range, 1994\u20132021)."
    )

    doc.add_heading('Overall Time Toxicity', level=2)
    doc.add_paragraph(
        "Median 12-month time toxicity was 19.0 days (IQR, 14.0\u201329.0) for intervention arms and "
        "18.0 days (IQR, 14.0\u201330.0) for control arms (Table 1). The corresponding time toxicity "
        "intensity was 1.6 days per month (IQR, 1.2\u20132.4) for intervention arms\u2014a burden "
        "comparable to routine off-protocol cancer management, in which a typical schedule of surveillance "
        "imaging, treatment administration, and clinic visits amounts to approximately 3 days per quarter. "
        "The mean (\u00b1 SD) time toxicity was 24.9 \u00b1 19.6 days for intervention arms and 24.4 "
        "\u00b1 20.3 days for control arms. Median incremental time toxicity (delta) was 0.0 days "
        "(IQR, 0.0\u20130.0); 431 trials (67.6%) imposed identical scheduling across both arms, "
        "reflecting shared Schedules of Assessments. Among the 207 trials with nonzero delta, the "
        "median incremental burden was 2.0 days (IQR, \u22125.0 to 9.0). A total of 128 trials (20.1%) "
        "showed higher time toxicity in the intervention arm, while 79 (12.4%) showed higher burden in "
        "the control arm. Among intervention arms, 144 (22.4%) exceeded 30 healthcare contact days, "
        "29 (4.5%) exceeded 60 days, and 8 (1.2%) exceeded 90 days."
    )

    doc.add_heading('Time Toxicity by Disease Site', level=2)
    doc.add_paragraph(
        "Time toxicity varied significantly across disease sites (Kruskal-Wallis H = 59.0, P < .001; "
        "Table 1, Figure 1). CNS trials had the highest median time toxicity (41.0 days; IQR, "
        "22.0\u201353.0; n = 13; intensity, 3.4 d/mo), followed by head and neck (25.0 days; IQR, "
        "16.0\u201339.0; n = 26) and gynecologic malignancies (24.0 days; IQR, 17.0\u201336.5; "
        "n = 40). Breast cancer trials had the lowest median among major disease sites (16.0 days; "
        "IQR, 8.5\u201320.5; n = 99; intensity, 1.3 d/mo), representing a 2.6-fold difference from "
        "CNS. Hematologic malignancies, despite comprising the largest disease category (134 [20.8%]), "
        "exhibited the widest variability (IQR, 14.0\u201342.0), reflecting substantial heterogeneity "
        "in protocol design across lymphoma, leukemia, and myeloma trials. Bonferroni-corrected "
        "pairwise comparisons identified 10 significant differences among disease sites. Thoracic "
        "tumors, despite moderate median time toxicity (18.0 days), had the narrowest variability "
        "(IQR, 17.0\u201325.5), suggesting greater consistency in protocol design within this disease site."
    )

    doc.add_heading('Time Toxicity by Treatment Modality', level=2)
    doc.add_paragraph(
        "Treatment modality was the strongest determinant of time toxicity (H = 105.6, P < .001; "
        "Table 1). Chemotherapy-based regimens were most burdensome (median, 32.0 days; IQR, "
        "20.0\u201352.0; n = 84; intensity, 2.7 d/mo), followed by immunotherapy (21.0 days; IQR, "
        "18.0\u201327.0; n = 137) and targeted therapy (19.0 days; IQR, 15.0\u201327.5; n = 283). "
        "Immunotherapy was significantly less burdensome than chemotherapy (P = .0002; rank-biserial "
        "r = 0.30)\u2014a finding of particular relevance given the ongoing shift in the treatment "
        "landscape from cytotoxic to immune-based regimens. Endocrine therapy had a median of 14.0 days "
        "(IQR, 7.8\u201317.0; n = 40), while other and supportive regimens had the lowest burden at "
        "11.0 days (IQR, 3.5\u201318.5; n = 83)."
    )

    doc.add_heading('Component Analysis', level=2)
    doc.add_paragraph(
        "Across 1431 arms with nonzero time toxicity, core treatment accounted for a mean of 68.6% of "
        "total healthcare contact days, followed by laboratory testing (16.7%), imaging and diagnostics "
        "(9.7%), and clinic visits (4.9%) (Figure 2). Research-specific monitoring (imaging plus "
        "laboratory testing) constituted 26.4% of total time toxicity overall\u2014representing a mean "
        "of 5.5 healthcare contact days per arm that could potentially be reduced through protocol "
        "optimization. This fraction varied substantially by disease site: hematologic trials had the "
        "highest research-specific burden (34.5%), driven by frequent laboratory monitoring (25.7% of "
        "total time toxicity), while gastrointestinal trials had the lowest (15.0%). By modality, "
        "immunotherapy had the lowest research-specific burden (7.1%), with 91.9% of time toxicity "
        "attributable to core treatment, whereas endocrine therapy trials allocated the highest "
        "proportion to research-specific monitoring (30.9%)."
    )

    # ========================================================================
    # DISCUSSION
    # ========================================================================
    doc.add_heading('Discussion', level=1)

    doc.add_paragraph(
        "The central finding of this analysis\u2014that median time toxicity intensity is 1.6 days "
        "per month\u2014reframes expectations about the burden of oncology trial participation. This "
        "intensity is comparable to what patients receiving standard-of-care cancer treatment experience "
        "off-protocol, where surveillance imaging, drug administration, and clinic visits typically "
        "amount to approximately 3 days per quarter. This challenges the prevailing assumption that "
        "clinical trial participation imposes a dramatically higher time burden than standard care and "
        "has important implications for how time toxicity is discussed with patients considering "
        "enrollment. Furthermore, the finding that 67.6% of trials impose identical scheduling on "
        "intervention and control arms\u2014because both arms share the same Schedule of "
        "Assessments\u2014indicates that the incremental time cost of experimental treatment is zero for "
        "most participants, with time toxicity driven by study design rather than treatment assignment. "
        "However, this reassuring aggregate figure masks clinically important variation: CNS trials "
        "impose an intensity of 3.4 days per month, and chemotherapy-based regimens 2.7 days per month, "
        "representing disease contexts where the time burden of trial participation warrants explicit "
        "discussion during the consent process."
    )

    doc.add_paragraph(
        "This 644-protocol dataset constitutes the first comprehensive reference standard for time "
        "toxicity across the phase III oncology trial landscape, providing what we believe to be a "
        "landmark resource for the field. The disease-site-specific and modality-specific benchmarks "
        "(Table 1) enable prospective comparison of proposed protocols against established norms for "
        "their disease context and treatment modality. The validated LLM extraction pipeline provides "
        "a practical, scalable instrument for time toxicity assessment\u2014processing full protocol "
        "documents with demonstrated reproducibility (>95% clinical stability)\u2014that could "
        "facilitate routine quantitative evaluation of protocol burden before trial activation. Time "
        "toxicity is also a potentially important consideration for clinical trial enrollment and "
        "retention: patients cite schedule burden as a leading barrier to participation, and quantifying "
        "this burden prospectively could improve informed consent and help sponsors design more "
        "patient-centered protocols."
    )

    doc.add_paragraph(
        "The finding that 26.4% of time toxicity is attributable to research-specific "
        "monitoring\u2014imaging and laboratory testing beyond core treatment requirements\u2014"
        "identifies a concrete, modifiable target. The principle of beneficence requires that the "
        "burdens of research participation be minimized and proportionate to potential benefits. "
        "Decentralized trial designs, remote monitoring, and streamlined surveillance protocols could "
        "meaningfully reduce this component, particularly in hematologic and gynecologic trials where "
        "research-specific burden exceeds 34%. These benchmark data provide an empirical foundation for "
        "evaluating whether proposed protocols impose burden proportionate to their disease context and "
        "for identifying specific protocols where schedule optimization could improve the participant "
        "experience."
    )

    doc.add_paragraph(
        "This study has several limitations. LLM extraction may introduce systematic bias in the "
        "interpretation of complex Schedule of Assessment tables, though the three-run consensus with "
        "median aggregation mitigates random error and demonstrated high stability. Protocol-mandated "
        "visit schedules may overestimate actual patient experience if visits are consolidated or "
        "skipped in practice; conversely, they may underestimate burden if unscheduled visits for "
        "toxicity management are required. Our analysis captures only the first 12 months and does "
        "not account for the burden of extended maintenance therapy or long-term follow-up surveillance. "
        "Component decomposition relied on two-run consensus owing to incomplete third-run category "
        "data. Finally, our dataset is limited to phase III trials registered on ClinicalTrials.gov "
        "and may not generalize to earlier-phase studies or trials conducted in low-resource settings."
    )

    # ========================================================================
    # TABLE 1
    # ========================================================================
    doc.add_page_break()
    add_para(doc, "Table 1. Time Toxicity by Disease Site and Treatment Modality (Intervention Arms)",
             bold=True, font_size=11, space_after=6)

    headers = ['', 'N', 'Median TT (d)', 'IQR', 'Intensity (d/mo)', 'Research (%)']

    disease_rows = [
        ['Panel A: By Disease Site', '', '', '', '', ''],
        ['CNS', '13', '41.0', '22.0\u201353.0', '3.4', '19.6'],
        ['Head & Neck', '26', '25.0', '16.0\u201339.0', '2.1', '11.3'],
        ['Gynecologic', '40', '24.0', '17.0\u201336.5', '2.0', '34.1'],
        ['Hematologic', '134', '22.0', '14.0\u201342.0', '1.8', '34.5'],
        ['Gastrointestinal', '80', '20.5', '16.8\u201328.2', '1.7', '15.0'],
        ['Skin', '29', '19.0', '16.0\u201325.0', '1.6', '19.2'],
        ['Thoracic', '94', '18.0', '17.0\u201325.5', '1.5', '18.5'],
        ['Genitourinary', '90', '18.0', '11.2\u201324.0', '1.5', '28.2'],
        ['Other', '28', '17.0', '10.0\u201349.2', '1.4', '28.6'],
        ['Breast', '99', '16.0', '8.5\u201320.5', '1.3', '20.6'],
        ['Multiple', '11', '6.0', '3.5\u201314.0', '0.5', '50.1'],
        ['', '', '', '', '', ''],
        ['Panel B: By Treatment Modality', '', '', '', '', ''],
        ['Chemotherapy', '84', '32.0', '20.0\u201352.0', '2.7', '26.9'],
        ['Radiation', '11', '27.0', '12.5\u201332.5', '2.3', '25.7'],
        ['Immunotherapy', '137', '21.0', '18.0\u201327.0', '1.8', '7.1'],
        ['Targeted Therapy', '283', '19.0', '15.0\u201327.5', '1.6', '29.0'],
        ['Endocrine', '40', '14.0', '7.8\u201317.0', '1.2', '30.9'],
        ['Other/Supportive', '83', '11.0', '3.5\u201318.5', '0.9', '34.5'],
    ]

    table = doc.add_table(rows=1 + len(disease_rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'

    # Data
    for r_idx, row_data in enumerate(disease_rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            if 'Panel' in val:
                run.bold = True
                run.italic = True

    doc.add_paragraph('')
    add_para(doc,
        "Abbreviations: TT, time toxicity (12-month healthcare contact days); IQR, interquartile range; "
        "Intensity, TT divided by 12 months; Research, percentage of TT attributable to research-specific "
        "monitoring (imaging + laboratory testing). "
        "Overall: median TT, 19.0 days (IQR, 14.0\u201329.0); intensity, 1.6 d/mo. "
        "Disease site comparison: Kruskal-Wallis H = 59.0, P < .001. "
        "Treatment modality comparison: H = 105.6, P < .001.",
        font_size=9, italic=True, space_after=12)

    # ========================================================================
    # FIGURE PLACEHOLDERS
    # ========================================================================
    doc.add_page_break()
    add_para(doc, "Figure 1. Distribution of 12-Month Time Toxicity by Disease Site (Intervention Arms)",
             bold=True, font_size=11, space_after=4)
    add_para(doc,
        "[Figure placeholder \u2014 box plot of 12-month time toxicity by disease site]",
        italic=True, space_after=4)
    doc.add_paragraph(
        "Horizontal box plots showing median and IQR of healthcare contact days for primary intervention "
        "arms, ordered by descending median. Vertical dashed line indicates overall median (19.0 days). "
        "Numbers in parentheses indicate sample size per group."
    )

    doc.add_paragraph('')
    add_para(doc, "Figure 2. Component Composition of Time Toxicity by Disease Site and Treatment Modality",
             bold=True, font_size=11, space_after=4)
    add_para(doc,
        "[Figure placeholder \u2014 stacked bar chart of TT components]",
        italic=True, space_after=4)
    doc.add_paragraph(
        "Stacked bar charts showing the proportion of 12-month healthcare contact days attributable to "
        "core treatment, laboratory testing, imaging and diagnostics, and clinic visits. Left panel: by "
        "disease site (intervention arms). Right panel: by treatment modality (intervention arms). "
        "Research-specific monitoring (imaging + labs) accounted for 26.4% of total time toxicity overall."
    )

    # ========================================================================
    # REFERENCES
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('References', level=1)

    refs = [
        "1. Gupta A, Eisenhauer EA, Booth CM. The time toxicity of cancer treatment. J Clin Oncol. 2022;40(15):1611-1615.",
        "2. Shah KN, Turchin A, Engel J, et al. Time toxicity: a framework for quantifying the patient burden of cancer treatment. J Natl Cancer Inst. 2024;116(3):345-352.",
        "3. Gupta A, Booth CM, Eisenhauer EA. Time toxicity of cancer treatment: a quality-of-life metric. Lancet Oncol. 2024;25(1):e1-e3.",
        "4. Gupta A, Stewart TL, Iqbal J, Booth CM. Time toxicity in oncology clinical trials: quantifying healthcare contact days. JAMA Oncol. 2024;10(8):1092-1099.",
        "5. ClinicalTrials.gov. US National Library of Medicine. Accessed January 2026. https://clinicaltrials.gov",
        "6. [Authors]. TimeToxSV: a large language model pipeline for automated extraction of healthcare contact days from clinical trial protocols. [Companion methods paper]. 2026.",
        "7. von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Ann Intern Med. 2007;147(8):573-577.",
        "8. Topol EJ. Artificial intelligence in health care: progress and challenges. N Engl J Med. 2023;388(3):258-263.",
        "9. Rajpurkar P, Chen E, Banerjee O, Topol EJ. AI in health and medicine. Nat Med. 2022;28(1):31-38.",
        "10. National Commission for the Protection of Human Subjects of Biomedical and Behavioral Research. The Belmont Report. Washington, DC: US DHEW; 1979.",
        "11. Khozin S, Blumenthal GM, Pazdur R. Real-world data for clinical evidence generation in oncology. J Natl Cancer Inst. 2017;109(11):djx187.",
        "12. Marra C, Chen JL, Coravos A, Stern AD. Quantifying the use of connected digital products in clinical research. NPJ Digit Med. 2020;3:50.",
        "13. Unger JM, Vaidya R, Hershman DL, Minasian LM, Fleury ME. Systematic review and meta-analysis of the magnitude of structural, clinical, and physician and patient barriers to cancer clinical trial participation. J Natl Cancer Inst. 2019;111(3):245-255.",
        "14. Winkfield KM, Phillips JK, Joffe S, et al. Addressing financial barriers to patient participation in clinical trials: ASCO policy statement. J Clin Oncol. 2018;36(33):3331-3339.",
        "15. Nipp RD, Hong K, Paskett ED. Overcoming barriers to clinical trial enrollment. Am Soc Clin Oncol Educ Book. 2019;39:105-114.",
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9)

    # ========================================================================
    # SAVE
    # ========================================================================
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, 'JAMA_Oncology_Brief_Report.docx')
    doc.save(output_path)
    print(f"Saved: {output_path}")
    print(f"Size: {os.path.getsize(output_path) / 1024:.0f} KB")


if __name__ == '__main__':
    main()
